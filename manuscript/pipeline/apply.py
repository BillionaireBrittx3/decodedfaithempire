import docx, json, sys, re
def smartify(s):
    # opening double quote after start/space/opening bracket, else closing
    s=re.sub(r'(^|[\s(\[{])"', r'\1“', s)
    s=s.replace('"','”')
    # opening single quote after start/space/opening bracket, else apostrophe/closing
    s=re.sub(r"(^|[\s(\[{])'", r"\1‘", s)
    s=s.replace("'", "’")
    return s
def set_text(p, new):
    new=smartify(new)
    if not p.runs:
        p.add_run(new); return
    p.runs[0].text = new
    for r in p.runs[1:]:
        r._element.getparent().remove(r._element)
def main():
    devos={r["n"]:r for r in json.load(open("devomap.json"))}
    rew=json.load(open("rewrites.json"))
    over=json.load(open("titles_override.json"))
    d=docx.Document("Man_source.docx"); P=d.paragraphs
    n_t=n_p=0
    for k,v in rew.items():
        r=devos[int(k)]
        if "teaching" in v:
            assert len(v["teaching"])==5, f"devo {k} teaching!=5"
            for idx,new in zip(r["teaching"], v["teaching"]):
                set_text(P[idx], new); n_t+=1
        if "prayer" in v:
            assert len(v["prayer"])==6, f"devo {k} prayer!=6"
            for idx,new in zip(r["prayer_body"], v["prayer"]):
                set_text(P[idx], new); n_p+=1
    for k,newtitle in over.items():
        r=devos[int(k)]
        set_text(P[r["title_i"]], newtitle)
        set_text(P[r["prayer_title_i"]], newtitle)
    out=sys.argv[1] if len(sys.argv)>1 else "Man_rebuilt.docx"
    d.save(out)
    print(f"applied {len(rew)} devotions | teaching paras={n_t} prayer paras={n_p} -> {out}")
if __name__=="__main__":
    main()
